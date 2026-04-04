#!/usr/bin/env python3
"""
Convert Kratom & Kava Directory JSON to PDF and DOCX with full contact information
"""

import json
import re
import os
from datetime import datetime

def extract_phone_from_address(address):
    """Extract phone number from address field"""
    if not address:
        return None
    # Match US phone formats: (XXX) XXX-XXXX or XXX-XXX-XXXX
    phone_pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    match = re.search(phone_pattern, address)
    if match:
        return match.group()
    return None

def extract_address_only(address):
    """Extract just the street address from the full address field"""
    if not address:
        return ""
    # Remove phone numbers
    cleaned = re.sub(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', '', address)
    # Remove rating, business type, hours, etc.
    lines = cleaned.split('\n')
    for line in lines:
        # Look for lines that look like addresses (contain numbers and street keywords)
        if any(keyword in line.lower() for keyword in ['blvd', 'ave', 'st', 'dr', 'rd', 'hwy', 'ln', 'ct', 'way', 'pl', 'unit', 'suite', '#']):
            return line.strip()
        # Or lines that start with numbers
        if re.match(r'^\d+\s', line.strip()):
            return line.strip()
    return ""

def load_json_data():
    """Load the JSON data file"""
    json_file = "kratom_kava_comprehensive_20260213_133443.json"
    if not os.path.exists(json_file):
        print(f"Error: {json_file} not found")
        return None
    
    with open(json_file, 'r', encoding='utf-8') as f:
        return json.load(f)

def create_txt_output(data):
    """Create a comprehensive text output with all contact info"""
    output = []
    output.append("=" * 80)
    output.append("KRATOM & KAVA DIRECTORY - SOUTH FLORIDA & NATIONWIDE")
    output.append(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    output.append("=" * 80)
    output.append("")
    
    # Kava Bars Section
    output.append("=" * 80)
    output.append("KAVA BARS")
    output.append("=" * 80)
    output.append("")
    
    # Group kava bars by location
    kava_bars = data.get('kava_bars', [])
    
    # Known database entries first
    output.append("-" * 40)
    output.append("MAJOR KAVA BAR CHAINS")
    output.append("-" * 40)
    for bar in kava_bars:
        if bar.get('source') == 'known_database' and bar.get('type') in ['kava_bar_chain', 'kava_bar']:
            output.append(f"\n{bar.get('name', 'Unknown')}")
            output.append(f"  Location: {bar.get('location', 'N/A')}")
            if bar.get('website'):
                output.append(f"  Website: {bar.get('website')}")
            output.append(f"  Type: {bar.get('type', 'Kava Bar')}")
    
    # Google Maps entries by location
    locations = {}
    for bar in kava_bars:
        if bar.get('source') == 'google_maps':
            loc = bar.get('location', 'Other')
            if loc not in locations:
                locations[loc] = []
            locations[loc].append(bar)
    
    for location, bars in sorted(locations.items()):
        output.append("")
        output.append("-" * 40)
        output.append(f"KAVA BARS - {location.upper()}")
        output.append("-" * 40)
        
        for bar in bars:
            output.append(f"\n{bar.get('name', 'Unknown')}")
            if bar.get('rating'):
                output.append(f"  Rating: {bar.get('rating')} stars")
            
            # Extract address
            address = extract_address_only(bar.get('address', ''))
            if address:
                output.append(f"  Address: {address}")
            
            # Extract phone
            phone = extract_phone_from_address(bar.get('address', ''))
            if phone:
                output.append(f"  Phone: {phone}")
            
            # Google Maps link
            if bar.get('link'):
                output.append(f"  Google Maps: {bar.get('link')}")
    
    # Kratom Vendors Section
    output.append("")
    output.append("=" * 80)
    output.append("KRATOM VENDORS")
    output.append("=" * 80)
    output.append("")
    
    kratom_vendors = data.get('kratom_vendors', [])
    
    # Online vendors
    output.append("-" * 40)
    output.append("ONLINE KRATOM VENDORS")
    output.append("-" * 40)
    for vendor in kratom_vendors:
        if vendor.get('type') == 'online_vendor' and vendor.get('source') == 'known_database':
            output.append(f"\n{vendor.get('name', 'Unknown')}")
            if vendor.get('website'):
                output.append(f"  Website: {vendor.get('website')}")
    
    # Brands
    output.append("")
    output.append("-" * 40)
    output.append("KRATOM BRANDS (Nationwide Distribution)")
    output.append("-" * 40)
    for vendor in kratom_vendors:
        if vendor.get('type') == 'brand':
            output.append(f"\n{vendor.get('name', 'Unknown')}")
            if vendor.get('website'):
                output.append(f"  Website: {vendor.get('website')}")
    
    # Local shops
    output.append("")
    output.append("-" * 40)
    output.append("LOCAL KRATOM SHOPS - SOUTH FLORIDA")
    output.append("-" * 40)
    for vendor in kratom_vendors:
        if vendor.get('source') == 'google_maps':
            output.append(f"\n{vendor.get('name', 'Unknown')}")
            if vendor.get('rating'):
                output.append(f"  Rating: {vendor.get('rating')} stars")
            
            address = extract_address_only(vendor.get('address', ''))
            if address:
                output.append(f"  Address: {address}")
            
            phone = extract_phone_from_address(vendor.get('address', ''))
            if phone:
                output.append(f"  Phone: {phone}")
            
            if vendor.get('link'):
                output.append(f"  Google Maps: {vendor.get('link')}")
    
    # Disclaimer
    output.append("")
    output.append("=" * 80)
    output.append("DISCLAIMER")
    output.append("=" * 80)
    output.append("")
    output.append("KRATOM LEGALITY: Kratom laws vary by state and locality. While kratom is legal")
    output.append("in Florida, it may be regulated or prohibited in other jurisdictions. Always")
    output.append("check local laws before purchasing or using kratom products.")
    output.append("")
    output.append("HEALTH NOTICE: The FDA has not approved kratom for any medical use. Consult a")
    output.append("healthcare professional before using kratom, especially if you have underlying")
    output.append("health conditions or are taking medications.")
    output.append("")
    output.append("DATA NOTE: Email addresses were not publicly available from the sources")
    output.append("scraped. Contact businesses directly via phone or website for email inquiries.")
    output.append("Some businesses listed may have permanently closed or relocated.")
    output.append("")
    output.append(f"Last Updated: {datetime.now().strftime('%B %d, %Y')}")
    
    return "\n".join(output)

def create_pdf_with_reportlab(data, filename):
    """Create PDF using reportlab with full contact info"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
    except ImportError:
        print("reportlab not installed. Install with: pip install reportlab")
        return False
    
    doc = SimpleDocTemplate(filename, pagesize=letter, 
                           rightMargin=0.5*inch, leftMargin=0.5*inch,
                           topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=colors.darkblue
    )
    
    section_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=15,
        spaceAfter=10,
        textColor=colors.darkblue,
        borderPadding=5,
        backColor=colors.lightgrey
    )
    
    subsection_style = ParagraphStyle(
        'SubsectionHeader',
        parent=styles['Heading3'],
        fontSize=12,
        spaceBefore=10,
        spaceAfter=5,
        textColor=colors.darkblue
    )
    
    business_style = ParagraphStyle(
        'BusinessName',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica-Bold',
        spaceBefore=8,
        spaceAfter=2
    )
    
    info_style = ParagraphStyle(
        'InfoStyle',
        parent=styles['Normal'],
        fontSize=9,
        leftIndent=15,
        spaceAfter=2
    )
    
    link_style = ParagraphStyle(
        'LinkStyle',
        parent=styles['Normal'],
        fontSize=9,
        leftIndent=15,
        textColor=colors.blue,
        spaceAfter=2
    )
    
    disclaimer_style = ParagraphStyle(
        'Disclaimer',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.darkred,
        spaceBefore=10,
        spaceAfter=5
    )
    
    story = []
    
    # Title
    story.append(Paragraph("KRATOM & KAVA DIRECTORY", title_style))
    story.append(Paragraph("South Florida & Nationwide", styles['Heading2']))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Statistics
    kava_count = len(data.get('kava_bars', []))
    kratom_count = len(data.get('kratom_vendors', []))
    stats = f"<b>Total Listings:</b> {kava_count} Kava Bars | {kratom_count} Kratom Vendors"
    story.append(Paragraph(stats, info_style))
    story.append(Spacer(1, 20))
    
    # Kava Bars Section
    story.append(Paragraph("KAVA BARS", section_style))
    
    kava_bars = data.get('kava_bars', [])
    
    # Major chains first
    story.append(Paragraph("Major Kava Bar Chains", subsection_style))
    for bar in kava_bars:
        if bar.get('source') == 'known_database' and bar.get('type') in ['kava_bar_chain', 'kava_bar']:
            story.append(Paragraph(bar.get('name', 'Unknown'), business_style))
            story.append(Paragraph(f"Location: {bar.get('location', 'N/A')}", info_style))
            if bar.get('website'):
                story.append(Paragraph(f"<link href='{bar.get('website')}'>Website: {bar.get('website')}</link>", link_style))
    
    # Google Maps entries by location
    locations = {}
    for bar in kava_bars:
        if bar.get('source') == 'google_maps':
            loc = bar.get('location', 'Other')
            if loc not in locations:
                locations[loc] = []
            locations[loc].append(bar)
    
    for location, bars in sorted(locations.items()):
        story.append(Paragraph(f"Kava Bars - {location}", subsection_style))
        
        for bar in bars[:10]:  # Limit to 10 per location for PDF size
            story.append(Paragraph(bar.get('name', 'Unknown'), business_style))
            
            if bar.get('rating'):
                story.append(Paragraph(f"Rating: {bar.get('rating')} stars", info_style))
            
            address = extract_address_only(bar.get('address', ''))
            if address:
                story.append(Paragraph(f"Address: {address}", info_style))
            
            phone = extract_phone_from_address(bar.get('address', ''))
            if phone:
                story.append(Paragraph(f"Phone: {phone}", info_style))
            
            if bar.get('link'):
                story.append(Paragraph(f"<link href='{bar.get('link')}'>Google Maps</link>", link_style))
    
    # Page break before kratom
    story.append(PageBreak())
    
    # Kratom Vendors Section
    story.append(Paragraph("KRATOM VENDORS", section_style))
    
    kratom_vendors = data.get('kratom_vendors', [])
    
    # Online vendors
    story.append(Paragraph("Online Kratom Vendors", subsection_style))
    for vendor in kratom_vendors:
        if vendor.get('type') == 'online_vendor' and vendor.get('source') == 'known_database':
            story.append(Paragraph(vendor.get('name', 'Unknown'), business_style))
            if vendor.get('website'):
                story.append(Paragraph(f"<link href='{vendor.get('website')}'>Website: {vendor.get('website')}</link>", link_style))
    
    # Brands
    story.append(Paragraph("Major Kratom Brands", subsection_style))
    for vendor in kratom_vendors:
        if vendor.get('type') == 'brand':
            story.append(Paragraph(vendor.get('name', 'Unknown'), business_style))
            if vendor.get('website'):
                story.append(Paragraph(f"<link href='{vendor.get('website')}'>Website: {vendor.get('website')}</link>", link_style))
    
    # Local shops
    story.append(Paragraph("Local Kratom Shops - South Florida", subsection_style))
    for vendor in kratom_vendors:
        if vendor.get('source') == 'google_maps':
            story.append(Paragraph(vendor.get('name', 'Unknown'), business_style))
            
            if vendor.get('rating'):
                story.append(Paragraph(f"Rating: {vendor.get('rating')} stars", info_style))
            
            address = extract_address_only(vendor.get('address', ''))
            if address:
                story.append(Paragraph(f"Address: {address}", info_style))
            
            phone = extract_phone_from_address(vendor.get('address', ''))
            if phone:
                story.append(Paragraph(f"Phone: {phone}", info_style))
            
            if vendor.get('link'):
                story.append(Paragraph(f"<link href='{vendor.get('link')}'>Google Maps</link>", link_style))
    
    # Disclaimer
    story.append(Spacer(1, 20))
    story.append(Paragraph("DISCLAIMER", section_style))
    story.append(Paragraph("<b>KRATOM LEGALITY:</b> Kratom laws vary by state and locality. While kratom is legal in Florida, it may be regulated or prohibited in other jurisdictions. Always check local laws before purchasing or using kratom products.", disclaimer_style))
    story.append(Paragraph("<b>HEALTH NOTICE:</b> The FDA has not approved kratom for any medical use. Consult a healthcare professional before using kratom.", disclaimer_style))
    story.append(Paragraph("<b>DATA NOTE:</b> Email addresses were not publicly available from the sources scraped. Contact businesses directly via phone or website for email inquiries.", disclaimer_style))
    
    doc.build(story)
    return True

def create_docx(data, filename):
    """Create DOCX using python-docx with full contact info"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("python-docx not installed. Install with: pip install python-docx")
        return False
    
    doc = Document()
    
    # Title
    title = doc.add_heading('KRATOM & KAVA DIRECTORY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('South Florida & Nationwide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Statistics
    kava_count = len(data.get('kava_bars', []))
    kratom_count = len(data.get('kratom_vendors', []))
    stats = doc.add_paragraph()
    stats.add_run(f'Total Listings: ').bold = True
    stats.add_run(f'{kava_count} Kava Bars | {kratom_count} Kratom Vendors')
    
    doc.add_paragraph()
    
    # Kava Bars Section
    doc.add_heading('KAVA BARS', level=1)
    
    kava_bars = data.get('kava_bars', [])
    
    # Major chains first
    doc.add_heading('Major Kava Bar Chains', level=2)
    for bar in kava_bars:
        if bar.get('source') == 'known_database' and bar.get('type') in ['kava_bar_chain', 'kava_bar']:
            p = doc.add_paragraph()
            p.add_run(bar.get('name', 'Unknown')).bold = True
            doc.add_paragraph(f'Location: {bar.get("location", "N/A")}', style='List Bullet')
            if bar.get('website'):
                doc.add_paragraph(f'Website: {bar.get("website")}', style='List Bullet')
    
    # Google Maps entries by location
    locations = {}
    for bar in kava_bars:
        if bar.get('source') == 'google_maps':
            loc = bar.get('location', 'Other')
            if loc not in locations:
                locations[loc] = []
            locations[loc].append(bar)
    
    for location, bars in sorted(locations.items()):
        doc.add_heading(f'Kava Bars - {location}', level=2)
        
        for bar in bars:
            p = doc.add_paragraph()
            p.add_run(bar.get('name', 'Unknown')).bold = True
            if bar.get('rating'):
                p.add_run(f' ({bar.get("rating")} stars)')
            
            address = extract_address_only(bar.get('address', ''))
            if address:
                doc.add_paragraph(f'Address: {address}', style='List Bullet')
            
            phone = extract_phone_from_address(bar.get('address', ''))
            if phone:
                doc.add_paragraph(f'Phone: {phone}', style='List Bullet')
            
            if bar.get('link'):
                doc.add_paragraph(f'Google Maps: {bar.get("link")}', style='List Bullet')
    
    # Page break
    doc.add_page_break()
    
    # Kratom Vendors Section
    doc.add_heading('KRATOM VENDORS', level=1)
    
    kratom_vendors = data.get('kratom_vendors', [])
    
    # Online vendors
    doc.add_heading('Online Kratom Vendors', level=2)
    for vendor in kratom_vendors:
        if vendor.get('type') == 'online_vendor' and vendor.get('source') == 'known_database':
            p = doc.add_paragraph()
            p.add_run(vendor.get('name', 'Unknown')).bold = True
            if vendor.get('website'):
                doc.add_paragraph(f'Website: {vendor.get("website")}', style='List Bullet')
    
    # Brands
    doc.add_heading('Major Kratom Brands', level=2)
    for vendor in kratom_vendors:
        if vendor.get('type') == 'brand':
            p = doc.add_paragraph()
            p.add_run(vendor.get('name', 'Unknown')).bold = True
            if vendor.get('website'):
                doc.add_paragraph(f'Website: {vendor.get("website")}', style='List Bullet')
    
    # Local shops
    doc.add_heading('Local Kratom Shops - South Florida', level=2)
    for vendor in kratom_vendors:
        if vendor.get('source') == 'google_maps':
            p = doc.add_paragraph()
            p.add_run(vendor.get('name', 'Unknown')).bold = True
            if vendor.get('rating'):
                p.add_run(f' ({vendor.get("rating")} stars)')
            
            address = extract_address_only(vendor.get('address', ''))
            if address:
                doc.add_paragraph(f'Address: {address}', style='List Bullet')
            
            phone = extract_phone_from_address(vendor.get('address', ''))
            if phone:
                doc.add_paragraph(f'Phone: {phone}', style='List Bullet')
            
            if vendor.get('link'):
                doc.add_paragraph(f'Google Maps: {vendor.get("link")}', style='List Bullet')
    
    # Disclaimer
    doc.add_heading('DISCLAIMER', level=1)
    
    disclaimer1 = doc.add_paragraph()
    disclaimer1.add_run('KRATOM LEGALITY: ').bold = True
    disclaimer1.add_run('Kratom laws vary by state and locality. While kratom is legal in Florida, it may be regulated or prohibited in other jurisdictions. Always check local laws before purchasing or using kratom products.')
    
    disclaimer2 = doc.add_paragraph()
    disclaimer2.add_run('HEALTH NOTICE: ').bold = True
    disclaimer2.add_run('The FDA has not approved kratom for any medical use. Consult a healthcare professional before using kratom, especially if you have underlying health conditions or are taking medications.')
    
    disclaimer3 = doc.add_paragraph()
    disclaimer3.add_run('DATA NOTE: ').bold = True
    disclaimer3.add_run('Email addresses were not publicly available from the sources scraped. Contact businesses directly via phone or website for email inquiries. Some businesses listed may have permanently closed or relocated.')
    
    doc.save(filename)
    return True

def main():
    print("Loading JSON data...")
    data = load_json_data()
    if not data:
        return
    
    print("Creating TXT output...")
    txt_content = create_txt_output(data)
    with open('kratom_kava_directory.txt', 'w', encoding='utf-8') as f:
        f.write(txt_content)
    print("Created: kratom_kava_directory.txt")
    
    print("Creating PDF...")
    if create_pdf_with_reportlab(data, 'kratom_kava_directory.pdf'):
        print("Created: kratom_kava_directory.pdf")
    
    print("Creating DOCX...")
    if create_docx(data, 'kratom_kava_directory.docx'):
        print("Created: kratom_kava_directory.docx")
    
    print("\nDone! All files created with full contact information including:")
    print("  - Business names")
    print("  - Addresses")
    print("  - Phone numbers (where available)")
    print("  - Website links (where available)")
    print("  - Google Maps links")

if __name__ == "__main__":
    main()
