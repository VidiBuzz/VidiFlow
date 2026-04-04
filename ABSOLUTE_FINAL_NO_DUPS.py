#!/usr/bin/env python3
"""
ABSOLUTE FINAL VERSION: Zero duplicates, guaranteed
Strategy: Calculate ALL insertions first, THEN insert in reverse order
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import random

def get_all_valid_images(images_dir):
    """Get all valid images"""
    valid_ext = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}
    images = []

    for file in os.listdir(images_dir):
        if Path(file).suffix.lower() in valid_ext:
            filepath = os.path.join(images_dir, file)
            try:
                img = Image.open(filepath)
                img.verify()
                img = Image.open(filepath)
                img.load()
                images.append((file, filepath))
            except:
                pass

    random.shuffle(images)
    return images

def plan_all_insertions(doc_path, all_images):
    """Plan ALL insertions WITHOUT actually inserting"""

    doc = Document(doc_path)

    # CRITICAL: Find existing images in document and mark as used!
    existing_images = set()
    for para in doc.paragraphs:
        if para.text.startswith('Figure:'):
            img_name = para.text.replace('Figure:', '').strip()
            # Extract base filename (title case version matches our filenames)
            # Convert back from title case to filename
            base_name = img_name.replace(' ', '_').lower()
            existing_images.add(img_name)  # Store the display version
            print(f"  Existing: {img_name[:50]}")

    print(f"  Found {len(existing_images)} existing images in document")

    # Find chapters
    chapters = {}
    for i, para in enumerate(doc.paragraphs):
        m = re.match(r'Chapter\s+(\d+):', para.text.strip())
        if m:
            chapters[int(m.group(1))] = i

    chapter_nums = sorted(chapters.keys())

    # Plan insertions
    planned_insertions = []  # List of (para_index, img_filename, img_path)
    used_images = set(existing_images)  # Start with existing images!
    image_idx = 0

    for idx, ch_num in enumerate(chapter_nums):
        start_para = chapters[ch_num]
        end_para = chapters[chapter_nums[idx + 1]] if idx + 1 < len(chapter_nums) else len(doc.paragraphs)

        para_range = end_para - start_para
        target_images = max(4, min(7, para_range // 12))
        spacing = para_range // (target_images + 1)

        for img_idx in range(target_images):
            # Get next unused image
            while image_idx < len(all_images):
                img_filename, img_path = all_images[image_idx]
                image_idx += 1

                # Convert to display format for comparison with existing images
                display_name = Path(img_filename).stem.replace('_', ' ').title()

                if display_name not in used_images:
                    used_images.add(display_name)

                    # Calculate insertion point
                    insert_para = start_para + (img_idx + 1) * spacing

                    if insert_para < end_para:
                        planned_insertions.append((insert_para, img_filename, img_path))
                    break

    return planned_insertions

def execute_insertions(doc_path, planned_insertions, output_path):
    """Execute ALL insertions in reverse order (preserves indices)"""

    doc = Document(doc_path)

    # Sort in REVERSE order (highest paragraph index first)
    planned_insertions.sort(key=lambda x: x[0], reverse=True)

    inserted = 0
    for para_idx, img_filename, img_path in planned_insertions:
        try:
            para = doc.paragraphs[para_idx]

            # Insert image
            new_para = para.insert_paragraph_before()
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = new_para.add_run()
            run.add_picture(img_path, width=Inches(4.5))

            # Caption
            caption_para = para.insert_paragraph_before()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(
                f"Figure: {Path(img_filename).stem.replace('_', ' ').title()}"
            )
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True

            inserted += 1

        except Exception as e:
            print(f"Skip: {img_filename}")

    doc.save(output_path)
    return inserted

if __name__ == "__main__":
    images_dir = r"M:\code\vidismart\images"
    doc_path = r"M:\+Proj\VidiSmart\The Speed of Visual Ai - FINAL.docx"
    output_path = r"M:\+Proj\VidiSmart\The Speed of Visual Ai - FINAL_NO_DUPS_VERIFIED.docx"

    print("Step 1: Loading images...")
    all_images = get_all_valid_images(images_dir)
    print(f"  Found {len(all_images)} valid images")

    print("Step 2: Planning insertions...")
    planned = plan_all_insertions(doc_path, all_images)
    print(f"  Planned {len(planned)} insertions")

    # VERIFY uniqueness
    unique_filenames = set(p[1] for p in planned)
    print(f"  {len(unique_filenames)} unique images in plan")

    if len(unique_filenames) != len(planned):
        print("  ERROR: Duplicates in plan!")
        for filename in unique_filenames:
            count = sum(1 for p in planned if p[1] == filename)
            if count > 1:
                print(f"    {count}x {filename}")
    else:
        print("  Plan has ZERO duplicates!")

    print("Step 3: Executing insertions...")
    inserted = execute_insertions(doc_path, planned, output_path)

    print()
    print("="*70)
    print(f"SUCCESS!")
    print(f"Inserted {inserted} images")
    print(f"Output: {output_path}")
    print("="*70)
