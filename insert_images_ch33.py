#!/usr/bin/env python3
"""
Simple script to insert images into Chapters 33-37 of the book
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

def get_valid_images(images_dir):
    """Get all valid image files from directory"""
    valid_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}
    images = []

    for file in os.listdir(images_dir):
        ext = Path(file).suffix.lower()
        if ext in valid_extensions:
            filepath = os.path.join(images_dir, file)
            try:
                # Test if image can be opened
                img = Image.open(filepath)
                img.verify()
                images.append((file, filepath))
            except:
                pass

    return images

def find_chapter_ranges(doc):
    """Find paragraph indices for chapters 33-37"""
    chapters = {}

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        match = re.match(r'Chapter\s+(\d+):', text)
        if match:
            chapter_num = int(match.group(1))
            if 33 <= chapter_num <= 37:
                chapters[chapter_num] = i

    # Add end marker (end of document or start of chapter 38)
    chapters[38] = len(doc.paragraphs)

    return chapters

def insert_images_in_chapters(doc_path, images_dir, output_path, start_ch=33, end_ch=37):
    """Insert images into specified chapters"""

    # Load document
    doc = Document(doc_path)

    # Get chapter ranges
    chapters = find_chapter_ranges(doc)

    print(f"Found chapters: {sorted([k for k in chapters.keys() if k < 38])}")

    # Get valid images
    images = get_valid_images(images_dir)
    print(f"Found {len(images)} valid images")

    # Calculate insertion points
    # We'll insert 1 image every N paragraphs in the chapter range
    start_para = chapters.get(start_ch)
    end_para = chapters.get(end_ch + 1, len(doc.paragraphs))

    if start_para is None:
        print(f"Error: Chapter {start_ch} not found")
        return 0

    print(f"Chapter {start_ch} starts at paragraph {start_para}")
    print(f"Chapter {end_ch} ends at paragraph {end_para}")

    para_range = end_para - start_para
    print(f"Range: {para_range} paragraphs")

    # Calculate spacing (insert every N paragraphs)
    num_images_to_insert = min(len(images), para_range // 15)  # ~1 image per 15 paragraphs
    spacing = para_range // max(num_images_to_insert, 1)

    print(f"Will insert ~{num_images_to_insert} images with spacing of ~{spacing} paragraphs")

    # Insert images (work backwards to preserve indices)
    images_inserted = 0
    used_indices = set()

    for img_idx in range(num_images_to_insert):
        # Calculate insertion point
        insert_offset = (img_idx + 1) * spacing
        para_idx = start_para + insert_offset

        # Make sure we're not past the end
        if para_idx >= end_para:
            break

        # Skip if already used
        if para_idx in used_indices:
            continue

        # Get image
        if img_idx >= len(images):
            break

        img_filename, img_path = images[img_idx]

        try:
            # Get paragraph
            para = doc.paragraphs[para_idx]

            # Create new paragraph for image
            new_para = para.insert_paragraph_before()
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add image
            run = new_para.add_run()
            run.add_picture(img_path, width=Inches(5.5))

            # Add caption
            caption_para = para.insert_paragraph_before()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(
                f"Figure: {Path(img_filename).stem.replace('_', ' ').title()}"
            )
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True

            images_inserted += 1
            used_indices.add(para_idx)
            print(f"  Inserted: {img_filename} at paragraph {para_idx}")

        except Exception as e:
            print(f"  Failed to insert {img_filename}: {e}")

    # Save document
    doc.save(output_path)
    print(f"\nSaved document with {images_inserted} images to: {output_path}")

    return images_inserted

if __name__ == "__main__":
    doc_path = r"M:\+Proj\VidiSmart\The Speed of Visual Ai - STEP1.docx"
    images_dir = r"M:\code\vidismart\images"
    output_path = r"M:\+Proj\VidiSmart\The Speed of Visual Ai - COMPLETE_WITH_IMAGES.docx"

    images_inserted = insert_images_in_chapters(
        doc_path,
        images_dir,
        output_path,
        start_ch=33,
        end_ch=37
    )

    print(f"\n{'='*60}")
    print(f"SUCCESS! Inserted {images_inserted} images")
    print(f"{'='*60}")
