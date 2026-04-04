#!/usr/bin/env python3
"""
Generate PDF and DOCX documents with complete contact information
Updated: ONE BIG GRAND TOTAL NUMBER at top
"""

import json
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load data
INPUT_FILE = "kratom_kava_complete_20260214_140628.json"

with open(INPUT_FILE, 'r', encoding='utf-8') as f:
    data = json.load(f)

kava_bars = data.get('kava_bars', [])
kratom_vendors = data.get('kratom_vendors', [])

# Calculate totals
total_businesses = len(kava_bars) + len(kratom_vendors)
total_with_phone = sum(1 for b in kava_bars if b.get('phone')) + sum(1 for b in kratom_vendors if b.get('phone'))
total_with_website = sum(1 for b in kava_bars if b.get('website')) + sum(1 for b in kratom_vendors if b.get('website'))

# Sort by location
def get_location_key(bar):
    loc = bar.get('location', '') or bar.get('scraped_address', '') or bar.get('address', '')
    if 'South Florida' in loc or 'Fort Lauderdale' in loc or 'Miami' in loc or 'Hollywood' in loc or 'Davie' in loc:
        return '0_' + loc
    return '1_' + loc

kava_bars_sorted = sorted(kava_bars, key=get_location_key)
kratom_vendors_sorted = sorted(kratom_vendors, key=get_location_key)

# Define yellow/gold color
YELLOW = colors.HexColor('#FFD700')
YELLOW_DARK = colors.HexColor('#FFA500')

# Generate PDF
print("Generating PDF...")
pdf_file = "kratom_kava_complete_directory.pdf"
doc = SimpleDocTemplate(pdf_file, pagesize=letter, 
                        rightMargin=0.5*inch, leftMargin=0.5*inch,
                        topMargin=0.5*inch, bottomMargin=0.5*inch)

styles = getSampleStyleSheet()
title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, 
                             textColor=YELLOW, alignment=TA_CENTER, spaceAfter=12)
heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14,
                               textColor=YELLOW, spaceAfter=10, spaceBefore=20)
normal_style = ParagraphStyle('Normal', parent=styles['Normal'], fontSize=10, leading=14)

story = []

# Title
story.append(Paragraph("Complete Kratom & Kava Directory", title_style))
story.append(Paragraph("South Florida & Beyond", ParagraphStyle('Subtitle', parent=styles['Normal'], 
              alignment=TA_CENTER, fontSize=12, textColor=colors.gray)))
story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", 
              ParagraphStyle('Date', parent=styles['Normal'], alignment=TA_CENTER, fontSize=10, 
              textColor=colors.gray, spaceAfter=30)))

# ONE BIG GRAND TOTAL BOX
grand_total_style = ParagraphStyle('GrandTotal', parent=styles['Normal'], 
                                   fontSize=48, alignment=TA_CENTER, 
                                   textColor=colors.HexColor('#1a1a2e'),
                                   fontName='Helvetica-Bold')

grand_label_style = ParagraphStyle('GrandLabel', parent=styles['Normal'],
                                   fontSize=14, alignment=TA_CENTER,
                                   textColor=colors.HexColor('#1a1a2e'),
                                   fontName='Helvetica-Bold',
                                   spaceAfter=5)

grand_desc_style = ParagraphStyle('GrandDesc', parent=styles['Normal'],
                                  fontSize=11, alignment=TA_CENTER,
                                  textColor=colors.HexColor('#333333'))

# Grand total table with yellow background
grand_box = Table([
    [Paragraph("TOTAL ADDRESSABLE MARKET", grand_label_style)],
    [Paragraph(str(total_businesses), grand_total_style)],
    [Paragraph(f"Total Kratom & Kava Businesses to Contact", grand_desc_style)],
], colWidths=[5*inch])
grand_box.setStyle(TableStyle([
    ('BACKGROUND', (0, 0), (-1, -1), YELLOW),
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ('TOPPADDING', (0, 0), (-1, -1), 20),
    ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
    ('LEFTPADDING', (0, 0), (-1, -1), 30),
    ('RIGHTPADDING', (0, 0), (-1, -1), 30),
    ('BOX', (0, 0), (-1, -1), 3, YELLOW_DARK),
]))

# Center the grand total box
outer_table = Table([[grand_box]], colWidths=[7*inch])
outer_table.setStyle(TableStyle([
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
]))
story.append(outer_table)
story.append(Spacer(1, 25))

# Breakdown section
story.append(Paragraph("Breakdown by Category", ParagraphStyle('BreakdownTitle', 
              parent=styles['Heading3'], fontSize=12, textColor=YELLOW, alignment=TA_CENTER, spaceAfter=15)))

breakdown_data = [
    [str(len(kava_bars)), 'Kava Bars', str(len(kratom_vendors)), 'Kratom Vendors'],
    [str(total_with_phone), 'Have Phone Numbers', str(total_with_website), 'Have Websites'],
    [str(sum(1 for b in kava_bars if b.get('phone'))), 'Kava Bars w/ Phone', 
     str(sum(1 for b in kava_bars if b.get('website'))), 'Kava Bars w/ Website'],
    [str(sum(1 for b in kratom_vendors if b.get('phone'))), 'Kratom Vendors w/ Phone',
     str(sum(1 for b in kratom_vendors if b.get('website'))), 'Kratom Vendors w/ Website'],
]

breakdown_table = Table(breakdown_data, colWidths=[0.8*inch, 1.6*inch, 0.8*inch, 1.6*inch])
breakdown_table.setStyle(TableStyle([
    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#FFF8DC')),
    ('TEXTCOLOR', (0, 0), (0, -1), YELLOW_DARK),
    ('TEXTCOLOR', (2, 0), (2, -1), YELLOW_DARK),
    ('TEXTCOLOR', (1, 0), (1, -1), colors.HexColor('#333333')),
    ('TEXTCOLOR', (3, 0), (3, -1), colors.HexColor('#333333')),
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
    ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
    ('FONTSIZE', (0, 0), (0, -1), 16),
    ('FONTSIZE', (2, 0), (2, -1), 16),
    ('FONTSIZE', (1, 0), (1, -1), 10),
    ('FONTSIZE', (3, 0), (3, -1), 10),
    ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
    ('TOPPADDING', (0, 0), (-1, -1), 10),
    ('GRID', (0, 0), (-1, -1), 1, YELLOW),
]))
story.append(breakdown_table)
story.append(Spacer(1, 25))

# Kava Bars section
story.append(Paragraph(f"KAVA BARS ({len(kava_bars)} locations)", heading_style))

for bar in kava_bars_sorted:
    name = bar.get('name', 'Unknown')
    address = bar.get('scraped_address') or bar.get('address') or bar.get('location', 'N/A')
    phone = bar.get('phone', 'N/A')
    website = bar.get('website', 'N/A')
    rating = bar.get('rating', '')
    maps_url = bar.get('google_maps_url') or bar.get('link', '')
    
    # Business card table
    card_data = [
        [Paragraph(f"<b>{name}</b>" + (f" ({rating} stars)" if rating else ""), normal_style)],
        [Paragraph(f"<b>Address:</b> {address}", normal_style)],
        [Paragraph(f"<b>Phone:</b> <link href='tel:{phone}'>{phone}</link>" if phone != 'N/A' else f"<b>Phone:</b> {phone}", 
                  ParagraphStyle('Phone', parent=normal_style, textColor=colors.HexColor('#4CAF50')) if phone != 'N/A' else normal_style)],
        [Paragraph(f"<b>Website:</b> <link href='{website}'>{website}</link>" if website != 'N/A' else f"<b>Website:</b> {website}",
                  ParagraphStyle('Web', parent=normal_style, textColor=colors.HexColor('#2196F3')) if website != 'N/A' else normal_style)],
    ]
    if maps_url:
        card_data.append([Paragraph(f"<b>Google Maps:</b> <link href='{maps_url}'>{maps_url}</link>", 
                        ParagraphStyle('Maps', parent=normal_style, textColor=colors.HexColor('#FF9800')))])
    
    card_table = Table(card_data, colWidths=[7*inch])
    card_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#fafafa')),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#e0e0e0')),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    story.append(card_table)
    story.append(Spacer(1, 8))

# Page break before kratom vendors
story.append(PageBreak())

# Kratom Vendors section
story.append(Paragraph(f"KRATOM VENDORS ({len(kratom_vendors)} locations)", heading_style))

for vendor in kratom_vendors_sorted:
    name = vendor.get('name', 'Unknown')
    address = vendor.get('scraped_address') or vendor.get('address') or vendor.get('location', 'N/A')
    phone = vendor.get('phone', 'N/A')
    website = vendor.get('website', 'N/A')
    rating = vendor.get('rating', '')
    maps_url = vendor.get('google_maps_url') or vendor.get('link', '')
    
    card_data = [
        [Paragraph(f"<b>{name}</b>" + (f" ({rating} stars)" if rating else ""), normal_style)],
        [Paragraph(f"<b>Address:</b> {address}", normal_style)],
        [Paragraph(f"<b>Phone:</b> <link href='tel:{phone}'>{phone}</link>" if phone != 'N/A' else f"<b>Phone:</b> {phone}",
                  ParagraphStyle('Phone', parent=normal_style, textColor=colors.HexColor('#4CAF50')) if phone != 'N/A' else normal_style)],
        [Paragraph(f"<b>Website:</b> <link href='{website}'>{website}</link>" if website != 'N/A' else f"<b>Website:</b> {website}",
                  ParagraphStyle('Web', parent=normal_style, textColor=colors.HexColor('#2196F3')) if website != 'N/A' else normal_style)],
    ]
    if maps_url:
        card_data.append([Paragraph(f"<b>Google Maps:</b> <link href='{maps_url}'>{maps_url}</link>",
                        ParagraphStyle('Maps', parent=normal_style, textColor=colors.HexColor('#FF9800')))])
    
    card_table = Table(card_data, colWidths=[7*inch])
    card_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#fafafa')),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#e0e0e0')),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    story.append(card_table)
    story.append(Spacer(1, 8))

# Footer note
story.append(Spacer(1, 20))
story.append(Paragraph("Note: Email addresses are typically not displayed on Google Business Profiles.", 
              ParagraphStyle('Note', parent=styles['Normal'], fontSize=9, textColor=colors.gray, alignment=TA_CENTER)))
story.append(Paragraph("Contact businesses directly via phone or website for email inquiries.",
              ParagraphStyle('Note', parent=styles['Normal'], fontSize=9, textColor=colors.gray, alignment=TA_CENTER)))

doc.build(story)
print(f"PDF created: {pdf_file}")

# Generate DOCX
print("Generating DOCX...")
docx_file = "kratom_kava_complete_directory.docx"
doc = Document()

# Title
title = doc.add_heading('Complete Kratom & Kava Directory', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(255, 215, 0)

subtitle = doc.add_paragraph('South Florida & Beyond')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ONE BIG GRAND TOTAL
doc.add_paragraph()
grand_box = doc.add_paragraph()
grand_box.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add the grand total with prominent formatting
label_run = grand_box.add_run('TOTAL ADDRESSABLE MARKET\n')
label_run.bold = True
label_run.font.size = Pt(14)
label_run.font.color.rgb = RGBColor(26, 26, 46)

number_run = grand_box.add_run(f'{total_businesses}\n')
number_run.bold = True
number_run.font.size = Pt(72)
number_run.font.color.rgb = RGBColor(26, 26, 46)

desc_run = grand_box.add_run('Total Kratom & Kava Businesses to Contact')
desc_run.font.size = Pt(12)
desc_run.font.color.rgb = RGBColor(51, 51, 51)

# Add shading to the paragraph (simulated with a table in DOCX)
doc.add_paragraph()

# Breakdown section
breakdown_heading = doc.add_heading('Breakdown by Category', level=2)
breakdown_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in breakdown_heading.runs:
    run.font.color.rgb = RGBColor(255, 215, 0)

breakdown = doc.add_paragraph()
breakdown.alignment = WD_ALIGN_PARAGRAPH.CENTER
breakdown.add_run(f'Kava Bars: {len(kava_bars)} | Kratom Vendors: {len(kratom_vendors)}\n').bold = True
breakdown.add_run(f'Have Phone Numbers: {total_with_phone} | Have Websites: {total_with_website}\n')
breakdown.add_run(f'Kava Bars w/ Phone: {sum(1 for b in kava_bars if b.get("phone"))} | Kava Bars w/ Website: {sum(1 for b in kava_bars if b.get("website"))}\n')
breakdown.add_run(f'Kratom Vendors w/ Phone: {sum(1 for b in kratom_vendors if b.get("phone"))} | Kratom Vendors w/ Website: {sum(1 for b in kratom_vendors if b.get("website"))}')

# Kava Bars
doc.add_page_break()
kava_heading = doc.add_heading(f'KAVA BARS ({len(kava_bars)} locations)', level=1)
for run in kava_heading.runs:
    run.font.color.rgb = RGBColor(255, 215, 0)

for bar in kava_bars_sorted:
    name = bar.get('name', 'Unknown')
    address = bar.get('scraped_address') or bar.get('address') or bar.get('location', 'N/A')
    phone = bar.get('phone', 'N/A')
    website = bar.get('website', 'N/A')
    rating = bar.get('rating', '')
    maps_url = bar.get('google_maps_url') or bar.get('link', '')
    
    heading = doc.add_heading(name, level=2)
    if rating:
        rating_run = heading.add_run(f' ({rating} stars)')
        rating_run.font.color.rgb = RGBColor(255, 215, 0)
    
    para = doc.add_paragraph()
    para.add_run('Address: ').bold = True
    para.add_run(f'{address}\n')
    
    para.add_run('Phone: ').bold = True
    if phone != 'N/A':
        phone_run = para.add_run(phone)
        phone_run.font.color.rgb = RGBColor(76, 175, 80)
    else:
        para.add_run(phone)
    para.add_run('\n')
    
    para.add_run('Website: ').bold = True
    if website != 'N/A':
        web_run = para.add_run(website)
        web_run.font.color.rgb = RGBColor(33, 150, 243)
    else:
        para.add_run(website)
    para.add_run('\n')
    
    if maps_url:
        para.add_run('Google Maps: ').bold = True
        maps_run = para.add_run(maps_url)
        maps_run.font.color.rgb = RGBColor(255, 152, 0)
    
    doc.add_paragraph()

# Kratom Vendors
doc.add_page_break()
kratom_heading = doc.add_heading(f'KRATOM VENDORS ({len(kratom_vendors)} locations)', level=1)
for run in kratom_heading.runs:
    run.font.color.rgb = RGBColor(255, 215, 0)

for vendor in kratom_vendors_sorted:
    name = vendor.get('name', 'Unknown')
    address = vendor.get('scraped_address') or vendor.get('address') or vendor.get('location', 'N/A')
    phone = vendor.get('phone', 'N/A')
    website = vendor.get('website', 'N/A')
    rating = vendor.get('rating', '')
    maps_url = vendor.get('google_maps_url') or bar.get('link', '')
    
    heading = doc.add_heading(name, level=2)
    if rating:
        rating_run = heading.add_run(f' ({rating} stars)')
        rating_run.font.color.rgb = RGBColor(255, 215, 0)
    
    para = doc.add_paragraph()
    para.add_run('Address: ').bold = True
    para.add_run(f'{address}\n')
    
    para.add_run('Phone: ').bold = True
    if phone != 'N/A':
        phone_run = para.add_run(phone)
        phone_run.font.color.rgb = RGBColor(76, 175, 80)
    else:
        para.add_run(phone)
    para.add_run('\n')
    
    para.add_run('Website: ').bold = True
    if website != 'N/A':
        web_run = para.add_run(website)
        web_run.font.color.rgb = RGBColor(33, 150, 243)
    else:
        para.add_run(website)
    para.add_run('\n')
    
    if maps_url:
        para.add_run('Google Maps: ').bold = True
        maps_run = para.add_run(maps_url)
        maps_run.font.color.rgb = RGBColor(255, 152, 0)
    
    doc.add_paragraph()

# Footer note
doc.add_paragraph()
note = doc.add_paragraph('Note: Email addresses are typically not displayed on Google Business Profiles. Contact businesses directly via phone or website for email inquiries.')
note.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(docx_file)
print(f"DOCX created: {docx_file}")

print("\nAll documents generated successfully!")