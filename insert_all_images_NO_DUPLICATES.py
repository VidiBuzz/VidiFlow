#!/usr/bin/env python3
"""
Insert ALL unique images into the book with NO DUPLICATES
Text wrapping enabled for better layout
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import random

def get_all_valid_images(images_dir):
    """Get ALL valid image files from directory"""
    valid_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}
    images = []

    for file in os.listdir(images_dir):
        ext = Path(file).suffix.lower()
        if ext in valid_extensions:
            filepath = os.path.join(images_dir, file)
            try:
                img = Image.open(filepath)
                img.verify()
                # Re-open after verify (verify closes it)
                img = Image.open(filepath)
                img.load()  # Actually load it to test
                images.append((file, filepath))
            except Exception as e:
                print(f"Skipping bad image: {file}")

    # Shuffle to randomize distribution
    random.shuffle(images)
    return images

def add_text_wrapping(run):
    """Add text wrapping to image - makes it flow with text"""
    try:
        # Get the inline shape (image)
        inline = run._element.xpath('.//wp:inline')
        if inline:
            inline_elem = inline[0]
            # Convert inline to anchor (floating)
            parent = inline_elem.getparent()
            anchor = OxmlElement('wp:anchor')

            # Copy attributes
            for key, value in inline_elem.attrib.items():
                anchor.set(key, value)

            # Set wrapping
            anchor.set('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}behindDoc', '0')

            # Add wrap square
            wrapSquare = OxmlElement('wp:wrapSquare')
            wrapSquare.set('wrapText', 'bothSides')
            anchor.append(wrapSquare)

            # Copy children from inline to anchor
            for child in inline_elem:
                anchor.append(child)

            # Replace inline with anchor
            parent.replace(inline_elem, anchor)
    except:
        pass  # If wrapping fails, image will just be inline (centered)

def insert_all_images(doc_path, images_dir, output_path, images_per_chapter=8):
    """Insert ALL unique images with NO duplicates"""

    # Load document
    doc = Document(doc_path)

    # Get ALL valid images
    all_images = get_all_valid_images(images_dir)
    print(f"Found {len(all_images)} total valid images")

    # Find all chapter boundaries
    chapters = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        match = re.match(r'Chapter\s+(\d+):', text)
        if match:
            chapter_num = int(match.group(1))
            chapters[chapter_num] = i

    chapter_nums = sorted(chapters.keys())
    print(f"Found {len(chapter_nums)} chapters: {chapter_nums}")

    # Track used images globally
    used_image_names = set()
    images_inserted = 0
    image_queue = list(all_images)  # Make a copy we can pop from

    # Process each chapter
    for idx, ch_num in enumerate(chapter_nums):
        start_para = chapters[ch_num]
        end_para = chapters.get(chapter_nums[idx + 1]) if idx + 1 < len(chapter_nums) else len(doc.paragraphs)

        para_range = end_para - start_para

        # Calculate how many images for this chapter (min 6, max based on chapter size)
        num_images = max(6, min(images_per_chapter, para_range // 10))

        # If we're running low on images, reduce count
        num_images = min(num_images, len(image_queue))

        if num_images == 0:
            print(f"Chapter {ch_num}: No more unique images available!")
            continue

        # Calculate spacing
        spacing = para_range // (num_images + 1)

        chapter_images = 0

        # Insert images in this chapter
        for img_idx in range(num_images):
            if not image_queue:
                break

            # Pop next unique image
            img_filename, img_path = image_queue.pop(0)

            # Skip if already used (shouldn't happen, but double-check)
            if img_filename in used_image_names:
                continue

            # Calculate insertion point
            insert_offset = (img_idx + 1) * spacing
            para_idx = start_para + insert_offset

            if para_idx >= end_para:
                # Put image back in queue
                image_queue.insert(0, (img_filename, img_path))
                break

            try:
                para = doc.paragraphs[para_idx]

                # Create new paragraph for image
                new_para = para.insert_paragraph_before()
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Add image with text wrapping
                run = new_para.add_run()

                # Use smaller width for text wrapping (4 inches instead of 5.5)
                run.add_picture(img_path, width=Inches(4.0))

                # Try to add text wrapping
                add_text_wrapping(run)

                # Add caption
                caption_para = para.insert_paragraph_before()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(
                    f"Figure: {Path(img_filename).stem.replace('_', ' ').title()}"
                )
                caption_run.font.size = Pt(9)
                caption_run.font.italic = True

                # Mark as used
                used_image_names.add(img_filename)
                images_inserted += 1
                chapter_images += 1

            except Exception as e:
                print(f"  SKIP ERROR: {img_filename} - {str(e)[:50]}")
                # Don't put it back - it's broken, skip it permanently
                used_image_names.add(img_filename)  # Mark as used so we don't try again

        print(f"Chapter {ch_num}: Inserted {chapter_images} unique images")

    # Save document
    doc.save(output_path)
    print(f"\n{'='*60}")
    print(f"SUCCESS!")
    print(f"Inserted {images_inserted} UNIQUE images (NO DUPLICATES)")
    print(f"Unused images remaining: {len(image_queue)}")
    print(f"Output: {output_path}")
    print(f"{'='*60}")

    return images_inserted

if __name__ == "__main__":
    doc_path = r"M:\+Proj\VidiSmart\The Speed of Visual Ai - FINAL.docx"
    images_dir = r"M:\code\vidismart\images"
    output_path = r"M:\+Proj\VidiSmart\The Speed of Visual Ai - FINAL_NO_DUPLICATES.docx"

    # Insert 8 images per chapter (will be 296 total for 37 chapters)
    images_inserted = insert_all_images(
        doc_path,
        images_dir,
        output_path,
        images_per_chapter=8
    )
