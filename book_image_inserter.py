#!/usr/bin/env python3
"""
Book Image Inserter - AI-powered image placement for Word documents
Analyzes document content and inserts images at appropriate locations using vision AI
"""

import os
import re
import sys
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import logging
from dataclasses import dataclass
from datetime import datetime

# Document processing
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# Image processing
try:
    from PIL import Image
except ImportError:
    print("Installing Pillow...")
    os.system(f"{sys.executable} -m pip install Pillow")
    from PIL import Image

# AI/ML for content analysis
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
except ImportError:
    print("Installing scikit-learn and numpy...")
    os.system(f"{sys.executable} -m pip install scikit-learn numpy")
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('image_inserter.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


@dataclass
class ImageInfo:
    """Information about an image file"""
    path: str
    filename: str
    keywords: List[str]
    category: str
    width: int
    height: int
    size_kb: float


@dataclass
class DocumentSection:
    """A section of the document with context"""
    heading: str
    content: str
    paragraph_index: int
    keywords: List[str]
    context_type: str  # 'introduction', 'technical', 'example', 'conclusion', etc.


class BookImageInserter:
    """Main class for intelligent image insertion into Word documents"""
    
    # Keywords that indicate good image insertion points
    IMAGE_CUE_PHRASES = [
        "figure", "illustration", "shown below", "see image", "as shown",
        "diagram", "visual", "chart", "graph", "screenshot", "example",
        "depicted", "represented", "visualize", "comparison", "architecture",
        "workflow", "pipeline", "framework", "overview", "structure"
    ]
    
    # Category mappings for images based on filename patterns
    CATEGORY_PATTERNS = {
        'ai_agents': ['agent', 'orchestrat', 'crewai', 'autogen', 'langchain'],
        'ai_models': ['model', 'llm', 'gpt', 'claude', 'gemini', 'llama'],
        'business': ['business', 'enterprise', 'company', 'corporate', 'industry'],
        'technology': ['tech', 'hardware', 'gpu', 'server', 'infrastructure'],
        'analytics': ['analytics', 'dashboard', 'metric', 'data', 'report'],
        'workflow': ['workflow', 'pipeline', 'process', 'flow', 'diagram'],
        'ui_screens': ['ui', 'interface', 'screenshot', 'app', 'screen'],
        'logos': ['logo', 'brand', 'icon'],
        'people': ['person', 'team', 'collaboration', 'meeting'],
        'construction': ['construction', 'building', 'architect', 'contractor'],
        'healthcare': ['health', 'medical', 'patient', 'clinical'],
        'logistics': ['logistics', 'fleet', 'truck', 'shipping', 'delivery'],
        'food': ['food', 'meal', 'grocery', 'kitchen', 'recipe'],
        'finance': ['finance', 'trading', 'payment', 'invoice', 'accounting'],
    }
    
    def __init__(self, images_dir: str, output_dir: str = None):
        self.images_dir = Path(images_dir)
        self.output_dir = Path(output_dir) if output_dir else self.images_dir.parent / "output"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        self.image_cache: Dict[str, ImageInfo] = {}
        self.document: Optional[Document] = None
        
        logger.info(f"Initialized BookImageInserter")
        logger.info(f"  Images directory: {self.images_dir}")
        logger.info(f"  Output directory: {self.output_dir}")
    
    def scan_images(self) -> List[ImageInfo]:
        """Scan the images directory and build image metadata"""
        logger.info(f"Scanning images in {self.images_dir}...")
        
        images = []
        valid_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.svg'}
        
        for img_path in self.images_dir.iterdir():
            if img_path.suffix.lower() not in valid_extensions:
                continue
            
            try:
                # Get image dimensions
                with Image.open(img_path) as img:
                    width, height = img.size
                
                # Extract keywords from filename
                keywords = self._extract_keywords_from_filename(img_path.name)
                
                # Determine category
                category = self._categorize_image(img_path.name, keywords)
                
                # Get file size
                size_kb = img_path.stat().st_size / 1024
                
                img_info = ImageInfo(
                    path=str(img_path),
                    filename=img_path.name,
                    keywords=keywords,
                    category=category,
                    width=width,
                    height=height,
                    size_kb=size_kb
                )
                
                images.append(img_info)
                self.image_cache[img_path.name] = img_info
                
            except Exception as e:
                logger.warning(f"Could not process image {img_path.name}: {e}")
        
        logger.info(f"Found {len(images)} valid images")
        return images
    
    def _extract_keywords_from_filename(self, filename: str) -> List[str]:
        """Extract meaningful keywords from a filename"""
        # Remove extension
        name = Path(filename).stem
        
        # Remove common suffixes (timestamps, etc.)
        name = re.sub(r'_\d{13,}$', '', name)
        name = re.sub(r'_\d{4,}$', '', name)
        
        # Split on common separators
        parts = re.split(r'[_\-\s,.]+', name.lower())
        
        # Filter out very short parts and common words
        stopwords = {'ai', 'the', 'and', 'or', 'of', 'in', 'to', 'for', 'with', 'on', 'at'}
        keywords = [p for p in parts if len(p) > 2 and p not in stopwords]
        
        return keywords
    
    def _categorize_image(self, filename: str, keywords: List[str]) -> str:
        """Categorize an image based on filename and keywords"""
        filename_lower = filename.lower()
        keywords_str = ' '.join(keywords)
        
        for category, patterns in self.CATEGORY_PATTERNS.items():
            for pattern in patterns:
                if pattern in filename_lower or pattern in keywords_str:
                    return category
        
        return 'general'
    
    def analyze_document(self, doc_path: str) -> List[DocumentSection]:
        """Analyze the Word document and extract sections with context"""
        logger.info(f"Analyzing document: {doc_path}")
        
        self.document = Document(doc_path)
        sections = []
        
        current_heading = "Introduction"
        current_content = []
        paragraph_index = 0
        
        for para in self.document.paragraphs:
            text = para.text.strip()
            
            if not text:
                # Empty paragraph - potential section break
                if current_content:
                    content_text = '\n'.join(current_content)
                    context_type = self._determine_context_type(current_heading, content_text)
                    keywords = self._extract_keywords_from_text(content_text)
                    
                    sections.append(DocumentSection(
                        heading=current_heading,
                        content=content_text,
                        paragraph_index=paragraph_index - len(current_content),
                        keywords=keywords,
                        context_type=context_type
                    ))
                    current_content = []
                continue
            
            # Check if this is a heading
            if para.style.name.startswith('Heading'):
                # Save previous section if exists
                if current_content:
                    content_text = '\n'.join(current_content)
                    context_type = self._determine_context_type(current_heading, content_text)
                    keywords = self._extract_keywords_from_text(content_text)
                    
                    sections.append(DocumentSection(
                        heading=current_heading,
                        content=content_text,
                        paragraph_index=paragraph_index - len(current_content),
                        keywords=keywords,
                        context_type=context_type
                    ))
                
                current_heading = text
                current_content = []
            else:
                current_content.append(text)
            
            paragraph_index += 1
        
        # Don't forget the last section
        if current_content:
            content_text = '\n'.join(current_content)
            context_type = self._determine_context_type(current_heading, content_text)
            keywords = self._extract_keywords_from_text(content_text)
            
            sections.append(DocumentSection(
                heading=current_heading,
                content=content_text,
                paragraph_index=paragraph_index - len(current_content),
                keywords=keywords,
                context_type=context_type
            ))
        
        logger.info(f"Extracted {len(sections)} document sections")
        return sections
    
    def _determine_context_type(self, heading: str, content: str) -> str:
        """Determine the type of content context"""
        heading_lower = heading.lower()
        content_lower = content.lower()
        
        if any(word in heading_lower for word in ['introduction', 'overview', 'beginning']):
            return 'introduction'
        elif any(word in heading_lower for word in ['conclusion', 'summary', 'final']):
            return 'conclusion'
        elif any(word in heading_lower for word in ['example', 'case study', 'success story']):
            return 'example'
        elif any(word in content_lower for word in ['technical', 'algorithm', 'implementation', 'code']):
            return 'technical'
        elif any(word in content_lower for word in ['figure', 'illustration', 'shown below', 'see image']):
            return 'image_cue'
        else:
            return 'general'
    
    def _extract_keywords_from_text(self, text: str) -> List[str]:
        """Extract meaningful keywords from text"""
        # Convert to lowercase and extract words
        words = re.findall(r'\b[a-z]{4,}\b', text.lower())
        
        # Remove common stopwords
        stopwords = {
            'that', 'this', 'with', 'from', 'have', 'will', 'been', 'were', 'their',
            'there', 'would', 'could', 'should', 'which', 'when', 'what', 'where',
            'than', 'then', 'also', 'just', 'more', 'some', 'into', 'over', 'such'
        }
        
        keywords = [w for w in words if w not in stopwords]
        
        # Get most frequent words
        from collections import Counter
        word_counts = Counter(keywords)
        top_keywords = [word for word, count in word_counts.most_common(10)]
        
        return top_keywords
    
    def find_image_insertion_points(
        self,
        sections: List[DocumentSection],
        images: List[ImageInfo],
        max_images_per_section: int = 1,
        min_score: float = 0.2
    ) -> List[Tuple[int, ImageInfo, str]]:
        """
        Find the best insertion points for images in the document
        Returns: List of (paragraph_index, image, reason)
        """
        logger.info("Finding optimal image insertion points...")

        insertions = []
        used_images = set()  # Track which images have been used to prevent duplicates

        for section in sections:
            # Find images that match this section's context
            matching_images = self._match_images_to_section(section, images, used_images, min_score)
            
            # Only insert 1 image per section maximum to avoid clustering
            num_images = min(1, len(matching_images))
            
            # Add the best matching image
            for i, (img, score, reason) in enumerate(matching_images[:num_images]):
                # Insert after the heading + some content
                insert_index = section.paragraph_index + min(3, len(section.content.split('\n')) // 2)
                insertions.append((insert_index, img, reason))
                used_images.add(img.path)  # Mark this image as used
        
        # Sort by insertion index (descending) so we insert from bottom to top
        insertions.sort(key=lambda x: x[0], reverse=True)
        
        logger.info(f"Found {len(insertions)} image insertion points (no duplicates)")
        return insertions
    
    def _match_images_to_section(
        self,
        section: DocumentSection,
        images: List[ImageInfo],
        used_images: set = None,
        min_score: float = 0.2
    ) -> List[Tuple[ImageInfo, float, str]]:
        """Match images to a document section based on content similarity"""

        if used_images is None:
            used_images = set()
        
        matches = []
        
        # Create text representation of section
        section_text = f"{section.heading} {section.content} {' '.join(section.keywords)}"
        
        for img in images:
            # Skip already used images to prevent duplicates
            if img.path in used_images:
                continue
                
            score = 0.0
            reasons = []
            
            # 1. Keyword matching
            keyword_matches = sum(1 for kw in img.keywords if kw in section_text.lower())
            if keyword_matches > 0:
                score += keyword_matches * 0.3
                reasons.append(f"keyword match ({keyword_matches})")
            
            # 2. Category matching
            for category, patterns in self.CATEGORY_PATTERNS.items():
                if img.category == category:
                    for pattern in patterns:
                        if pattern in section_text.lower():
                            score += 0.4
                            reasons.append(f"category match ({category})")
                            break
            
            # 3. Cue phrase bonus
            if section.context_type == 'image_cue':
                score += 0.3
                reasons.append("cue phrase detected")
            
            # 4. Image quality bonus (prefer medium-sized images)
            if 400 <= img.width <= 1200 and 300 <= img.height <= 900:
                score += 0.1
                reasons.append("good dimensions")
            
            # 5. Penalize very small or very large images
            if img.width < 200 or img.height < 200:
                score -= 0.3
                reasons.append("too small")
            elif img.width > 2000 or img.height > 2000:
                score -= 0.2
                reasons.append("too large")
            
            # 6. Penalize duplicate filenames (same base name with different timestamps)
            base_name = re.sub(r'_\d{10,}', '', img.filename.lower())
            if any(re.sub(r'_\d{10,}', '', used_img.split('/')[-1].lower()) == base_name 
                   for used_img in used_images):
                score -= 0.5
                reasons.append("similar image already used")
            
            if score > min_score:  # Minimum threshold
                reason = f"Score: {score:.2f} - {', '.join(reasons)}"
                matches.append((img, score, reason))
        
        # Sort by score (descending) and take only the best match
        matches.sort(key=lambda x: x[1], reverse=True)
        
        # Return only the single best match to prevent clustering
        return matches[:1] if matches else []
    
    def insert_images(
        self,
        doc_path: str,
        output_path: str,
        image_width: float = 5.0,
        start_chapter: int = None,
        end_chapter: int = None
    ) -> int:
        """Insert images into the document at determined positions"""
        logger.info(f"Inserting images into document...")

        # Load document
        self.document = Document(doc_path)

        # Find chapter boundaries if filtering is requested
        chapter_start_para = 0
        chapter_end_para = len(self.document.paragraphs)

        if start_chapter is not None or end_chapter is not None:
            logger.info(f"Filtering to chapters {start_chapter} - {end_chapter}")
            for i, para in enumerate(self.document.paragraphs):
                text = para.text.strip()
                # Match "Chapter XX:" pattern
                import re
                chapter_match = re.match(r'Chapter\s+(\d+):', text)
                if chapter_match:
                    chapter_num = int(chapter_match.group(1))
                    if start_chapter is not None and chapter_num == start_chapter:
                        chapter_start_para = i
                        logger.info(f"Found Chapter {chapter_num} start at paragraph {i}")
                    if end_chapter is not None and chapter_num == end_chapter + 1:
                        chapter_end_para = i
                        logger.info(f"Found Chapter {chapter_num} (end boundary) at paragraph {i}")
                        break

        # Scan images
        images = self.scan_images()

        # Analyze document
        sections = self.analyze_document(doc_path)

        # Use lower threshold for targeted chapter ranges
        min_score = 0.0 if (start_chapter is not None or end_chapter is not None) else 0.2

        # Find insertion points
        insertions = self.find_image_insertion_points(sections, images, min_score=min_score)

        # Filter insertions to chapter range
        if start_chapter is not None or end_chapter is not None:
            original_count = len(insertions)
            insertions = [(idx, img, reason) for idx, img, reason in insertions
                         if chapter_start_para <= idx < chapter_end_para]
            logger.info(f"Filtered insertions from {original_count} to {len(insertions)} based on chapter range")
        
        # Insert images
        images_inserted = 0
        
        for para_index, img_info, reason in insertions:
            try:
                # Get the paragraph to insert before
                if para_index < len(self.document.paragraphs):
                    para = self.document.paragraphs[para_index]
                    
                    # Create a new paragraph for the image (insert before target)
                    new_para = para.insert_paragraph_before()
                    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add image with appropriate width
                    run = new_para.add_run()
                    
                    # Calculate width (constrain to reasonable size)
                    width_inches = min(image_width, 6.0)
                    run.add_picture(img_info.path, width=Inches(width_inches))
                    
                    # Add caption as a separate paragraph
                    caption_para = para.insert_paragraph_before()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(f"Figure: {Path(img_info.filename).stem.replace('_', ' ').title()}")
                    caption_run.font.size = Pt(9)
                    caption_run.font.italic = True
                    
                    images_inserted += 1
                    logger.info(f"  Inserted: {img_info.filename} before paragraph {para_index} - {reason}")
                    
            except Exception as e:
                logger.error(f"  Failed to insert {img_info.filename}: {e}")
        
        # Save document
        self.document.save(output_path)
        logger.info(f"Saved document with {images_inserted} images to: {output_path}")
        
        return images_inserted
    
    def generate_report(self, doc_path: str, output_path: str) -> str:
        """Generate a report of the analysis without modifying the document"""
        logger.info(f"Generating analysis report...")
        
        # Scan images
        images = self.scan_images()
        
        # Analyze document
        sections = self.analyze_document(doc_path)
        
        # Find insertion points
        insertions = self.find_image_insertion_points(sections, images)
        
        # Generate report
        report_lines = [
            "=" * 80,
            "BOOK IMAGE INSERTION ANALYSIS REPORT",
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "=" * 80,
            "",
            f"Document: {doc_path}",
            f"Images Available: {len(images)}",
            f"Document Sections: {len(sections)}",
            f"Recommended Insertions: {len(insertions)}",
            "",
            "-" * 80,
            "AVAILABLE IMAGES BY CATEGORY",
            "-" * 80,
        ]
        
        # Group images by category
        images_by_category = {}
        for img in images:
            if img.category not in images_by_category:
                images_by_category[img.category] = []
            images_by_category[img.category].append(img)
        
        for category, imgs in sorted(images_by_category.items()):
            report_lines.append(f"\n{category.upper()} ({len(imgs)} images):")
            for img in imgs[:5]:  # Show first 5
                report_lines.append(f"  - {img.filename} ({img.width}x{img.height}, {img.size_kb:.1f}KB)")
            if len(imgs) > 5:
                report_lines.append(f"  ... and {len(imgs) - 5} more")
        
        report_lines.extend([
            "",
            "-" * 80,
            "DOCUMENT STRUCTURE",
            "-" * 80,
        ])
        
        for i, section in enumerate(sections):
            report_lines.extend([
                f"\nSection {i+1}: {section.heading}",
                f"  Type: {section.context_type}",
                f"  Content Length: {len(section.content)} characters",
                f"  Keywords: {', '.join(section.keywords[:5])}",
            ])
        
        report_lines.extend([
            "",
            "-" * 80,
            "RECOMMENDED IMAGE PLACEMENTS",
            "-" * 80,
        ])
        
        for para_index, img_info, reason in insertions:
            # Find which section this belongs to
            section_name = "Unknown"
            for section in sections:
                if para_index >= section.paragraph_index:
                    section_name = section.heading
                    break
            
            report_lines.extend([
                f"\nAfter paragraph {para_index} (in section: {section_name}):",
                f"  Image: {img_info.filename}",
                f"  Reason: {reason}",
            ])
        
        report_lines.extend([
            "",
            "=" * 80,
            "END OF REPORT",
            "=" * 80,
        ])
        
        report_text = '\n'.join(report_lines)
        
        # Save report
        report_path = self.output_dir / "analysis_report.txt"
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(report_text)
        
        logger.info(f"Report saved to: {report_path}")
        return report_text


def main():
    """Main entry point"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="AI-powered image insertion for Word documents"
    )
    parser.add_argument(
        "document",
        help="Path to the Word document to process"
    )
    parser.add_argument(
        "--images-dir",
        default="M:\\code\\vidismart\\images",
        help="Directory containing images to insert"
    )
    parser.add_argument(
        "--output",
        help="Output path for the modified document (default: adds _with_images suffix)"
    )
    parser.add_argument(
        "--report-only",
        action="store_true",
        help="Generate analysis report without modifying document"
    )
    parser.add_argument(
        "--image-width",
        type=float,
        default=5.0,
        help="Width of inserted images in inches (default: 5.0)"
    )
    parser.add_argument(
        "--start-chapter",
        type=int,
        help="Start inserting images from this chapter number (inclusive)"
    )
    parser.add_argument(
        "--end-chapter",
        type=int,
        help="Stop inserting images at this chapter number (inclusive)"
    )

    args = parser.parse_args()
    
    # Validate inputs
    if not os.path.exists(args.document):
        print(f"Error: Document not found: {args.document}")
        sys.exit(1)
    
    if not os.path.exists(args.images_dir):
        print(f"Error: Images directory not found: {args.images_dir}")
        sys.exit(1)
    
    # Determine output path
    if args.output:
        output_path = args.output
    else:
        doc_path = Path(args.document)
        output_path = str(doc_path.parent / f"{doc_path.stem}_with_images{doc_path.suffix}")
    
    # Create inserter
    inserter = BookImageInserter(args.images_dir)
    
    if args.report_only:
        # Generate report only
        report = inserter.generate_report(args.document, output_path)
        print("\n" + report)
    else:
        # Insert images
        print(f"\nProcessing document: {args.document}")
        print(f"Images directory: {args.images_dir}")
        print(f"Output path: {output_path}")
        print()
        
        images_inserted = inserter.insert_images(
            args.document,
            output_path,
            image_width=args.image_width,
            start_chapter=args.start_chapter,
            end_chapter=args.end_chapter
        )
        
        print(f"\n{'=' * 60}")
        print(f"SUCCESS!")
        print(f"Inserted {images_inserted} images into the document")
        print(f"Output saved to: {output_path}")
        print(f"{'=' * 60}")


if __name__ == "__main__":
    main()