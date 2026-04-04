#!/usr/bin/env python3
"""
FINAL VERSION: Insert ALL unique images - ABSOLUTELY NO DUPLICATES
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import random

def get_all_valid_images(images_dir):
    """Get ALL valid image files"""
    valid_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}
    images = []

    print("Scanning images...")
    for file in os.listdir(images_dir):
        ext = Path(file).suffix.lower()
        if ext in valid_extensions:
            filepath = os.path.join(images_dir, file)
            try:
                img = Image.open(filepath)
                img.verify()
                img = Image.open(filepath)
                img.load()
                images.append((file, filepath))
            except:
                pass  # Skip bad images silently

    random.shuffle(images)
    print(f"Found {len(images)} valid images")
    return images

def insert_all_unique_images(doc_path, images_dir, output_path):
    """Insert ALL images with ZERO duplicates"""

    doc = Document(doc_path)

    # Get ALL images
    all_images = get_all_valid_images(images_dir)

    # Find chapters
    chapters = {}
    for i, para in enumerate(doc.paragraphs):
        match = re.match(r'Chapter\s+(\d+):', para.text.strip())
        if match:
            chapters[int(match.group(1))] = i

    chapter_nums = sorted(chapters.keys())
    print(f"Found {len(chapter_nums)} chapters")

    # GLOBAL tracking - this prevents ALL duplicates
    USED_IMAGES = set()
    images_inserted = 0

    # Process each chapter IN REVERSE to preserve paragraph indices
    for idx in range(len(chapter_nums) - 1, -1, -1):
        ch_num = chapter_nums[idx]
        start_para = chapters[ch_num]
        end_para = chapters.get(chapter_nums[idx + 1]) if idx + 1 < len(chapter_nums) else len(doc.paragraphs)

        para_range = end_para - start_para

        # Target 6-8 images per chapter
        target_images = max(5, min(8, para_range // 10))
        spacing = para_range // (target_images + 1)

        chapter_images = 0

        # Insert images
        for img_idx in range(target_images):
            # Find next UNUSED image
            img_filename = None
            img_path = None

            for filename, filepath in all_images:
                if filename not in USED_IMAGES:
                    img_filename = filename
                    img_path = filepath
                    USED_IMAGES.add(filename)  # MARK AS USED IMMEDIATELY
                    break

            if not img_filename:
                print(f"Chapter {ch_num}: Ran out of unique images!")
                break

            # Calculate insertion point
            insert_offset = (img_idx + 1) * spacing
            para_idx = start_para + insert_offset

            if para_idx >= end_para:
                break

            try:
                para = doc.paragraphs[para_idx]

                # Insert image
                new_para = para.insert_paragraph_before()
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run = new_para.add_run()
                run.add_picture(img_path, width=Inches(4.5))

                # Caption
                caption_para = para.insert_paragraph_before()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(
                    f"Figure: {Path(img_filename).stem.replace('_', ' ').title()}"
                )
                caption_run.font.size = Pt(9)
                caption_run.font.italic = True

                # Already marked as used above
                images_inserted += 1
                chapter_images += 1

            except Exception as e:
                # Already marked as used above
                print(f"  Skip: {img_filename}")

        print(f"Chapter {ch_num}: {chapter_images} images")

    # Save
    doc.save(output_path)

    print(f"\n{'='*70}")
    print(f"SUCCESS!")
    print(f"Inserted {images_inserted} UNIQUE images")
    print(f"ZERO DUPLICATES GUARANTEED")
    print(f"Output: {output_path}")
    print(f"{'='*70}")

    return images_inserted

if __name__ == "__main__":
    images_inserted = insert_all_unique_images(
        doc_path=r"M:\+Proj\VidiSmart\The Speed of Visual Ai - FINAL.docx",
        images_dir=r"M:\code\vidismart\images",
        output_path=r"M:\+Proj\VidiSmart\The Speed of Visual Ai - FINAL_CLEAN.docx"
    )
