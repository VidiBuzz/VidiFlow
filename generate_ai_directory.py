#!/usr/bin/env python3
"""
Generate AI & Visual AI Consulting Company Directory
HTML, TXT, PDF, DOCX outputs with complete contact information
"""

import json
import os
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load data
INPUT_FILE = "ai_visual_rag_companies_20260214_165715.json"

with open(INPUT_FILE, 'r', encoding='utf-8') as f:
    data = json.load(f)

businesses = data.get('businesses', [])

# Calculate totals
total_businesses = len(businesses)
total_with_phone = sum(1 for b in businesses if b.get('phone'))
total_with_website = sum(1 for b in businesses if b.get('website'))

# Group by region
regions = {}
for b in businesses:
    region = b.get('region', 'Unknown')
    if region not in regions:
        regions[region] = []
    regions[region].append(b)

# Sort regions
region_order = ['South Florida', 'San Francisco CA', 'Denver CO', 'Detroit MI', 'Chicago IL', 'Saint Louis MO', 'Cleveland OH', 'Bentonville AR']

# Define yellow/gold color
YELLOW = colors.HexColor('#FFD700')
YELLOW_DARK = colors.HexColor('#FFA500')

# Generate HTML
print("Generating HTML...")
html_file = "ai_visual_rag_directory.html"

html = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>AI & Visual AI Consulting Directory - Complete Company Listings</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
            min-height: 100vh;
            color: #e0e0e0;
            line-height: 1.6;
        }
        .container { max-width: 1400px; margin: 0 auto; padding: 20px; }
        header {
            text-align: center;
            padding: 40px 20px;
            background: rgba(0,0,0,0.3);
            border-bottom: 3px solid #FFD700;
            margin-bottom: 30px;
        }
        h1 { 
            color: #FFD700;
            font-size: 2.5em;
            margin-bottom: 10px;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.5);
        }
        .subtitle { color: #a0a0a0; font-size: 1.1em; margin-bottom: 20px; }
        
        /* ONE BIG GRAND TOTAL */
        .grand-total-box {
            background: linear-gradient(135deg, #FFD700 0%, #FFA500 100%);
            border-radius: 20px;
            padding: 40px;
            margin: 30px auto;
            max-width: 600px;
            text-align: center;
            box-shadow: 0 10px 40px rgba(255, 215, 0, 0.3);
        }
        .grand-total-label {
            color: #1a1a2e;
            font-size: 1.3em;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 2px;
            margin-bottom: 10px;
        }
        .grand-total-number {
            color: #1a1a2e;
            font-size: 6em;
            font-weight: 900;
            line-height: 1;
        }
        .grand-total-desc {
            color: #333;
            font-size: 1.1em;
            margin-top: 10px;
        }
        
        /* Breakdown section */
        .breakdown {
            background: rgba(0,0,0,0.2);
            border-radius: 15px;
            padding: 25px;
            margin: 20px auto;
            max-width: 900px;
            border: 1px solid rgba(255, 215, 0, 0.3);
        }
        .breakdown h3 {
            color: #FFD700;
            text-align: center;
            margin-bottom: 20px;
            font-size: 1.2em;
        }
        .breakdown-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
            gap: 15px;
        }
        .breakdown-item {
            background: rgba(255,255,255,0.05);
            padding: 15px;
            border-radius: 10px;
            text-align: center;
        }
        .breakdown-number {
            font-size: 1.8em;
            color: #FFD700;
            font-weight: bold;
        }
        .breakdown-label {
            color: #aaa;
            font-size: 0.85em;
            margin-top: 5px;
        }
        
        nav {
            position: sticky;
            top: 0;
            background: rgba(22, 33, 62, 0.95);
            padding: 15px;
            border-radius: 10px;
            margin-bottom: 20px;
            z-index: 100;
            display: flex;
            justify-content: center;
            gap: 15px;
            flex-wrap: wrap;
        }
        nav a {
            color: #FFD700;
            text-decoration: none;
            padding: 8px 15px;
            border: 1px solid #FFD700;
            border-radius: 5px;
            transition: all 0.3s;
            font-size: 0.9em;
        }
        nav a:hover { background: #FFD700; color: #1a1a2e; }
        
        .section {
            background: rgba(0,0,0,0.2);
            border-radius: 15px;
            padding: 25px;
            margin-bottom: 30px;
            border: 1px solid rgba(255, 215, 0, 0.3);
        }
        .section h2 {
            color: #FFD700;
            border-bottom: 2px solid #FFD700;
            padding-bottom: 10px;
            margin-bottom: 20px;
        }
        
        .business-grid {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(350px, 1fr));
            gap: 20px;
        }
        .business-card {
            background: rgba(255,255,255,0.05);
            border-radius: 10px;
            padding: 20px;
            border: 1px solid rgba(255,255,255,0.1);
            transition: all 0.3s;
        }
        .business-card:hover {
            border-color: #FFD700;
            transform: translateY(-3px);
            box-shadow: 0 5px 20px rgba(255, 215, 0, 0.2);
        }
        .business-name {
            font-size: 1.2em;
            color: #fff;
            margin-bottom: 10px;
            font-weight: 600;
        }
        .business-category {
            display: inline-block;
            background: rgba(255, 215, 0, 0.2);
            color: #FFD700;
            padding: 3px 10px;
            border-radius: 15px;
            font-size: 0.75em;
            margin-left: 10px;
        }
        .business-address {
            color: #a0a0a0;
            font-size: 0.9em;
            margin-bottom: 10px;
        }
        .contact-info {
            margin-top: 15px;
        }
        .contact-item {
            display: flex;
            align-items: center;
            margin-bottom: 8px;
            gap: 10px;
        }
        .contact-label {
            color: #888;
            min-width: 70px;
            font-size: 0.85em;
        }
        a.phone-link {
            color: #4CAF50;
            text-decoration: none;
            font-weight: 500;
        }
        a.phone-link:hover { text-decoration: underline; }
        a.website-link {
            color: #2196F3;
            text-decoration: none;
            word-break: break-all;
        }
        a.website-link:hover { text-decoration: underline; }
        a.maps-link {
            color: #FF9800;
            text-decoration: none;
            font-size: 0.9em;
        }
        a.maps-link:hover { text-decoration: underline; }
        
        .rating {
            display: inline-block;
            background: #FFD700;
            color: #1a1a2e;
            padding: 3px 10px;
            border-radius: 15px;
            font-size: 0.85em;
            margin-left: 10px;
            font-weight: bold;
        }
        .social-links {
            margin-top: 10px;
            padding-top: 10px;
            border-top: 1px solid rgba(255,255,255,0.1);
        }
        .social-link {
            display: inline-block;
            background: rgba(255,255,255,0.1);
            padding: 5px 12px;
            border-radius: 15px;
            margin-right: 8px;
            margin-bottom: 5px;
            font-size: 0.85em;
            color: #ccc;
            text-decoration: none;
        }
        .social-link:hover { background: rgba(255,255,255,0.2); }
        
        .missing {
            color: #666;
            font-style: italic;
        }
        
        footer {
            text-align: center;
            padding: 30px;
            color: #666;
            border-top: 1px solid rgba(255,255,255,0.1);
            margin-top: 40px;
        }
        
        @media (max-width: 768px) {
            .business-grid { grid-template-columns: 1fr; }
            h1 { font-size: 1.8em; }
            .grand-total-number { font-size: 4em; }
            .breakdown-grid { grid-template-columns: repeat(2, 1fr); }
        }
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>AI & Visual AI Consulting Directory</h1>
            <p class="subtitle">Complete Company Listings - AI, RAG, Visual AI, On-Premise Solutions</p>
            <p style="color: #888; font-size: 0.9em;">Generated: ''' + datetime.now().strftime('%B %d, %Y at %I:%M %p') + '''</p>
        </header>
        
        <!-- ONE BIG GRAND TOTAL -->
        <div class="grand-total-box">
            <div class="grand-total-label">TOTAL ADDRESSABLE MARKET</div>
            <div class="grand-total-number">''' + str(total_businesses) + '''</div>
            <div class="grand-total-desc">AI & Visual AI Consulting Companies to Contact</div>
        </div>
        
        <!-- BREAKDOWN -->
        <div class="breakdown">
            <h3>Breakdown by Region</h3>
            <div class="breakdown-grid">
'''

# Add region breakdown
for region in region_order:
    if region in regions:
        count = len(regions[region])
        html += f'''
                <div class="breakdown-item">
                    <div class="breakdown-number">{count}</div>
                    <div class="breakdown-label">{region}</div>
                </div>'''

html += f'''
                <div class="breakdown-item">
                    <div class="breakdown-number">{total_with_phone}</div>
                    <div class="breakdown-label">Have Phone Numbers</div>
                </div>
                <div class="breakdown-item">
                    <div class="breakdown-number">{total_with_website}</div>
                    <div class="breakdown-label">Have Websites</div>
                </div>
            </div>
        </div>
        
        <nav>
'''

# Add nav links
for region in region_order:
    if region in regions:
        html += f'            <a href="#{region.lower().replace(" ", "-")}">{region} ({len(regions[region])})</a>\n'

html += '''        </nav>
'''

# Add sections for each region
for region in region_order:
    if region not in regions:
        continue
    
    region_businesses = regions[region]
    html += f'''
        <section id="{region.lower().replace(" ", "-")}" class="section">
            <h2>{region} ({len(region_businesses)} companies)</h2>
            <div class="business-grid">
'''
    
    for biz in region_businesses:
        name = biz.get('name', 'Unknown')
        address = biz.get('scraped_address') or biz.get('address') or biz.get('location', 'Address not available')
        phone = biz.get('phone')
        website = biz.get('website')
        rating = biz.get('rating')
        category = biz.get('category', '')
        maps_url = biz.get('google_maps_url') or biz.get('link', '')
        social_links = biz.get('social_links', [])
        
        html += f'''
                <div class="business-card">
                    <div class="business-name">{name}'''
        if category:
            html += f'<span class="business-category">{category}</span>'
        if rating:
            html += f'<span class="rating">{"{:.1f}".format(float(rating)) if rating else ""} stars</span>'
        html += '''</div>
                    <div class="business-address">''' + str(address) + '''</div>
                    <div class="contact-info">
                        <div class="contact-item">
                            <span class="contact-label">Phone:</span>'''
        if phone:
            html += f'<a href="tel:{phone}" class="phone-link">{phone}</a>'
        else:
            html += '<span class="missing">Not available</span>'
        html += '''</div>
                        <div class="contact-item">
                            <span class="contact-label">Website:</span>'''
        if website:
            html += f'<a href="{website}" target="_blank" class="website-link">{website}</a>'
        else:
            html += '<span class="missing">Not available</span>'
        html += '''</div>'''
        if maps_url:
            html += f'''
                        <div class="contact-item">
                            <span class="contact-label">Maps:</span>
                            <a href="{maps_url}" target="_blank" class="maps-link">View on Google Maps</a>
                        </div>'''
        if social_links:
            html += '''
                        <div class="social-links">'''
            for social in social_links[:3]:
                html += f'<a href="{social["url"]}" target="_blank" class="social-link">{social["platform"]}</a>'
            html += '''</div>'''
        html += '''
                    </div>
                </div>'''
    
    html += '''
            </div>
        </section>'''

html += '''
        <footer>
            <p>This directory was automatically generated from Google Maps business listings.</p>
            <p>Companies specialize in: AI Consulting, Visual AI, RAG/LLM, On-Premise AI, Video Commerce</p>
            <p>Contact businesses directly via phone or website for email inquiries.</p>
        </footer>
    </div>
</body>
</html>
'''

with open(html_file, 'w', encoding='utf-8') as f:
    f.write(html)

print(f"HTML created: {html_file}")

# Generate TXT
print("Generating TXT...")
txt_file = "ai_visual_rag_directory.txt"

txt = f"""
{'='*80}
AI & VISUAL AI CONSULTING DIRECTORY
Complete Company Listings
Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
{'='*80}

{'='*80}
TOTAL ADDRESSABLE MARKET: {total_businesses} COMPANIES
{'='*80}

If you contacted every company on this list, you would be contacting {total_businesses} AI & Visual AI consulting companies.

BREAKDOWN BY REGION:
{'='*80}
"""

for region in region_order:
    if region in regions:
        txt += f"{region}: {len(regions[region])}\n"

txt += f"""
Have Phone Numbers: {total_with_phone}
Have Websites: {total_with_website}

"""

for region in region_order:
    if region not in regions:
        continue
    
    region_businesses = regions[region]
    txt += f"""
{'='*80}
{region.upper()} ({len(region_businesses)} companies)
{'='*80}
"""
    
    for biz in region_businesses:
        name = biz.get('name', 'Unknown')
        address = biz.get('scraped_address') or biz.get('address') or biz.get('location', 'N/A')
        phone = biz.get('phone', 'N/A')
        website = biz.get('website', 'N/A')
        rating = biz.get('rating', 'N/A')
        category = biz.get('category', 'N/A')
        maps_url = biz.get('google_maps_url') or biz.get('link', 'N/A')
        
        txt += f"""
{name}
{'-'*60}
Category: {category}
Address: {address}
Phone: {phone}
Website: {website}
Rating: {rating} stars
Google Maps: {maps_url}
"""

with open(txt_file, 'w', encoding='utf-8') as f:
    f.write(txt)

print(f"TXT created: {txt_file}")

# Generate PDF
print("Generating PDF...")
pdf_file = "ai_visual_rag_directory.pdf"
doc = SimpleDocTemplate(pdf_file, pagesize=letter, 
                        rightMargin=0.5*inch, leftMargin=0.5*inch,
                        topMargin=0.5*inch, bottomMargin=0.5*inch)

styles = getSampleStyleSheet()
title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, 
                             textColor=YELLOW, alignment=TA_CENTER, spaceAfter=12)
heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14,
                               textColor=YELLOW, spaceAfter=10, spaceBefore=20)
normal_style = ParagraphStyle('Normal', parent=styles['Normal'], fontSize=10, leading=14)

story = []

# Title
story.append(Paragraph("AI & Visual AI Consulting Directory", title_style))
story.append(Paragraph("Complete Company Listings", ParagraphStyle('Subtitle', parent=styles['Normal'], 
              alignment=TA_CENTER, fontSize=12, textColor=colors.gray)))
story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", 
              ParagraphStyle('Date', parent=styles['Normal'], alignment=TA_CENTER, fontSize=10, 
              textColor=colors.gray, spaceAfter=30)))

# Grand Total Box
grand_total_style = ParagraphStyle('GrandTotal', parent=styles['Normal'], 
                                   fontSize=48, alignment=TA_CENTER, 
                                   textColor=colors.HexColor('#1a1a2e'),
                                   fontName='Helvetica-Bold')
grand_label_style = ParagraphStyle('GrandLabel', parent=styles['Normal'],
                                   fontSize=14, alignment=TA_CENTER,
                                   textColor=colors.HexColor('#1a1a2e'),
                                   fontName='Helvetica-Bold',
                                   spaceAfter=5)
grand_desc_style = ParagraphStyle('GrandDesc', parent=styles['Normal'],
                                  fontSize=11, alignment=TA_CENTER,
                                  textColor=colors.HexColor('#333333'))

grand_box = Table([
    [Paragraph("TOTAL ADDRESSABLE MARKET", grand_label_style)],
    [Paragraph(str(total_businesses), grand_total_style)],
    [Paragraph(f"AI & Visual AI Consulting Companies to Contact", grand_desc_style)],
], colWidths=[5*inch])
grand_box.setStyle(TableStyle([
    ('BACKGROUND', (0, 0), (-1, -1), YELLOW),
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ('TOPPADDING', (0, 0), (-1, -1), 20),
    ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
    ('BOX', (0, 0), (-1, -1), 3, YELLOW_DARK),
]))

outer_table = Table([[grand_box]], colWidths=[7*inch])
outer_table.setStyle(TableStyle([('ALIGN', (0, 0), (-1, -1), 'CENTER')]))
story.append(outer_table)
story.append(Spacer(1, 25))

# Region breakdown
breakdown_data = []
for region in region_order:
    if region in regions:
        breakdown_data.append([str(len(regions[region])), region])
breakdown_data.append([str(total_with_phone), 'Have Phone Numbers'])
breakdown_data.append([str(total_with_website), 'Have Websites'])

# Split into rows of 4
breakdown_rows = []
for i in range(0, len(breakdown_data), 4):
    row = []
    for j in range(4):
        if i + j < len(breakdown_data):
            row.extend(breakdown_data[i + j])
        else:
            row.extend(['', ''])
    breakdown_rows.append(row)

breakdown_table = Table(breakdown_rows, colWidths=[0.6*inch, 1.4*inch, 0.6*inch, 1.4*inch])
breakdown_table.setStyle(TableStyle([
    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#FFF8DC')),
    ('TEXTCOLOR', (0, 0), (0, -1), YELLOW_DARK),
    ('TEXTCOLOR', (2, 0), (2, -1), YELLOW_DARK),
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
    ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
    ('FONTSIZE', (0, 0), (0, -1), 14),
    ('FONTSIZE', (2, 0), (2, -1), 14),
    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ('TOPPADDING', (0, 0), (-1, -1), 8),
    ('GRID', (0, 0), (-1, -1), 1, YELLOW),
]))
story.append(breakdown_table)
story.append(Spacer(1, 25))

# Add businesses by region
for region in region_order:
    if region not in regions:
        continue
    
    region_businesses = regions[region]
    story.append(Paragraph(f"{region} ({len(region_businesses)} companies)", heading_style))
    
    for biz in region_businesses:
        name = biz.get('name', 'Unknown')
        address = biz.get('scraped_address') or biz.get('address') or biz.get('location', 'N/A')
        phone = biz.get('phone', 'N/A')
        website = biz.get('website', 'N/A')
        rating = biz.get('rating', '')
        category = biz.get('category', '')
        maps_url = biz.get('google_maps_url') or biz.get('link', '')
        
        card_data = [
            [Paragraph(f"<b>{name}</b>" + (f" ({category})" if category else ""), normal_style)],
            [Paragraph(f"<b>Address:</b> {address}", normal_style)],
            [Paragraph(f"<b>Phone:</b> {phone}", ParagraphStyle('Phone', parent=normal_style, 
                       textColor=colors.HexColor('#4CAF50')) if phone != 'N/A' else normal_style)],
            [Paragraph(f"<b>Website:</b> {website}", ParagraphStyle('Web', parent=normal_style,
                       textColor=colors.HexColor('#2196F3')) if website != 'N/A' else normal_style)],
        ]
        if maps_url:
            card_data.append([Paragraph(f"<b>Google Maps:</b> {maps_url}", 
                            ParagraphStyle('Maps', parent=normal_style, textColor=colors.HexColor('#FF9800')))])
        
        card_table = Table(card_data, colWidths=[7*inch])
        card_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#fafafa')),
            ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#e0e0e0')),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(card_table)
        story.append(Spacer(1, 6))
    
    story.append(PageBreak())

# Footer
story.append(Paragraph("Note: Email addresses are typically not displayed on Google Business Profiles.", 
              ParagraphStyle('Note', parent=styles['Normal'], fontSize=9, textColor=colors.gray, alignment=TA_CENTER)))

doc.build(story)
print(f"PDF created: {pdf_file}")

# Generate DOCX
print("Generating DOCX...")
docx_file = "ai_visual_rag_directory.docx"
doc = Document()

# Title
title = doc.add_heading('AI & Visual AI Consulting Directory', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(255, 215, 0)

subtitle = doc.add_paragraph('Complete Company Listings')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Grand Total
doc.add_paragraph()
grand_box = doc.add_paragraph()
grand_box.alignment = WD_ALIGN_PARAGRAPH.CENTER

label_run = grand_box.add_run('TOTAL ADDRESSABLE MARKET\n')
label_run.bold = True
label_run.font.size = Pt(14)
label_run.font.color.rgb = RGBColor(26, 26, 46)

number_run = grand_box.add_run(f'{total_businesses}\n')
number_run.bold = True
number_run.font.size = Pt(72)
number_run.font.color.rgb = RGBColor(26, 26, 46)

desc_run = grand_box.add_run('AI & Visual AI Consulting Companies to Contact')
desc_run.font.size = Pt(12)
desc_run.font.color.rgb = RGBColor(51, 51, 51)

# Breakdown
doc.add_paragraph()
breakdown = doc.add_paragraph()
breakdown.alignment = WD_ALIGN_PARAGRAPH.CENTER
for region in region_order:
    if region in regions:
        breakdown.add_run(f'{region}: {len(regions[region])} | ').bold = True
breakdown.add_run(f'\nHave Phone Numbers: {total_with_phone} | Have Websites: {total_with_website}')

# Businesses by region
for region in region_order:
    if region not in regions:
        continue
    
    doc.add_page_break()
    region_heading = doc.add_heading(f'{region} ({len(regions[region])} companies)', level=1)
    for run in region_heading.runs:
        run.font.color.rgb = RGBColor(255, 215, 0)
    
    for biz in regions[region]:
        name = biz.get('name', 'Unknown')
        address = biz.get('scraped_address') or biz.get('address') or biz.get('location', 'N/A')
        phone = biz.get('phone', 'N/A')
        website = biz.get('website', 'N/A')
        category = biz.get('category', '')
        
        heading = doc.add_heading(name, level=2)
        if category:
            cat_run = heading.add_run(f' ({category})')
            cat_run.font.color.rgb = RGBColor(255, 215, 0)
        
        para = doc.add_paragraph()
        para.add_run('Address: ').bold = True
        para.add_run(f'{address}\n')
        
        para.add_run('Phone: ').bold = True
        if phone != 'N/A':
            phone_run = para.add_run(phone)
            phone_run.font.color.rgb = RGBColor(76, 175, 80)
        else:
            para.add_run(phone)
        para.add_run('\n')
        
        para.add_run('Website: ').bold = True
        if website != 'N/A':
            web_run = para.add_run(website)
            web_run.font.color.rgb = RGBColor(33, 150, 243)
        else:
            para.add_run(website)
        
        doc.add_paragraph()

doc.save(docx_file)
print(f"DOCX created: {docx_file}")

print("\nAll documents generated successfully!")
print(f"\nSummary:")
print(f"  Total Companies: {total_businesses}")
print(f"  With Phone: {total_with_phone}")
print(f"  With Website: {total_with_website}")